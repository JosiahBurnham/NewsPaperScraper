from bs4 import BeautifulSoup
import docx
from datetime import date
from docx.shared import Inches
import pathlib
import pandas as pd
import requests
import sys

from requests.models import MissingSchema
sys.path.insert(1, str(pathlib.Path(__file__).parent.absolute()))
import hyperlink


class MakePaper:
    """
    This class is a web scraper that makes a newspaper from newsites that have all the necessary
    html information in excel spreadsheets that are read using pandas.
    """

    def __init__(self):
        """initializes the document, data frames, and sets the title."""
        self.document = docx.Document()
        self.article_df = pd.read_excel(str(pathlib.Path(__file__).parent.absolute()) + "\\News_Paper_Links.xls")
        self.standalone_images_df = pd.read_excel(
            str(pathlib.Path(__file__).parent.absolute()) + "\\Standalone_Image_Links.xls")

        # adds a title to the news paper
        self.document.add_heading('YOUR TITLE HERE', 0)

    @staticmethod
    def get_image(link, picture_name):
        """ gets downloads and image from a link
        
        Uses requests to download the images from the link.
        Then saves the pictures to the Images folder in the same directory.
        
        :param link: the link to the picture to download it.
        :param picture_name: the name of the file to save to in Images directory.
        :return: none
        """
        response = requests.get(link)
        image = open(str(pathlib.Path(__file__).parent.absolute()) + '\\Images\\' + picture_name, 'wb')
        image.write(response.content)
        image.close()


    def get_date(self):
        """ Gets today's date
        
        :return: a string of today's date in the standard US format (MM/DD/YYYY)
        """
        self.document.add_paragraph(date.today().strftime('%B %d %Y'), style='Heading 1')

    def save_standalone_images(self):
        """ save images that are a feature by themselves
        
        Gets the image links and titles that that are standalone in the news paper from the excel file.
        It then uses the links and titles in get_image().
        
        :return: none
        """
        image_counter = 0
        # gets the links from the excel file and uses
        for links in self.standalone_images_df["URL"]:
            self.get_image(links, self.standalone_images_df["image_name"][image_counter])
            image_counter += 1

    def add_standalone_images(self):
        """ adds images that are their own feature to the document.
        
        Adds all of the standalone pictures to the document like the daily forecast, and tide chart.
        
        :return: none
        """
        for image_names in self.standalone_images_df["image_name"]:
            # adds the pictures to the document
            self.document.add_picture(str(pathlib.Path(__file__).parent.absolute()) + '\\Images\\' + image_names
                                      , width=Inches(5.0))

    def add_articles(self):
        """ adds all the articles to the document
        
        goes through each news site in the News_Paper_links excel file and get the article titles,
         images, and first paragraph, and then adds them to the document.
         
        :return: none
        """
        article_counter = 0

        # goes through each URL of news sites that are in the News_Paper_Links excel file
        for news_site in self.article_df["URL"]:

            source =requests.get(news_site)
            soup = BeautifulSoup(source.text, "lxml")

            # find where the articles are on the website
            main_page = soup.find(class_=self.article_df["article_location"][article_counter])
            for article in main_page.select(self.article_df["article_html_tag"][article_counter]):
                try:
                    # get title from main page
                    article_title = article.find(class_=self.article_df["article_title_html_tag"][article_counter])
                    self.document.add_paragraph(article_title.a.text, style='Heading 1')

                    # gets article source
                    href = article_title.a.get('href')
                    article_source = requests.get(href).text
                    article_soup = BeautifulSoup(article_source, 'lxml')

                    # gets article image
                    article_image_class = article_soup.find(class_=self.article_df["article_image_html_tag"][article_counter])
                    self.get_image(article_image_class.a.img.get('src'), "article_image.jpg")
                    self.document.add_picture(str(pathlib.Path(__file__).parent.absolute()) + '\\Images\\article_image.jpg',width=Inches(5.0))

                    # get first paragraph
                    article_text = article_soup.find(class_=self.article_df["article_text_html_tag"][article_counter])
                    starting_paragraph = article_text.find(class_=self.article_df["first_paragraph_html_tag"][article_counter])

                    # adds first paragraph to document
                    self.document.add_paragraph('Summary: ' + starting_paragraph.text)

                    try:
                        # adds a hyper link to the article if available
                        p = self.document.add_paragraph('')
                        hyperlink.Hyperlink().add_hyperlink(p, 'Article', href)

                    except UnboundLocalError:
                        print('', end='')
                except (AttributeError, MissingSchema):
                    print('', end='')

    def make_paper(self):
        """ makes the document and saves it.
        
        This just combines all the previous methods to add all the information to the document.
        It then saves the document.
        
        :return: none
        """
        self.get_date()
        self.save_standalone_images()
        self.add_standalone_images()
        self.add_articles()
        self.document.save(str(pathlib.Path(__file__).parent.absolute()) + '\\YOUR_TITLE_HERE.docx')



